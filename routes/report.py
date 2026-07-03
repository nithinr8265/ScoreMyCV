import json
import io
from flask import Blueprint, request, jsonify, session, send_file, make_response
from utils.supabase_client import get_supabase

report_bp = Blueprint("report", __name__)

def login_required(fn):
    from functools import wraps
    @wraps(fn)
    def wrapper(*args, **kwargs):
        if not session.get("user"):
            return jsonify({"error": "Authentication required"}), 401
        return fn(*args, **kwargs)
    return wrapper

def _get_analysis(analysis_id, user_id):
    sb = get_supabase()
    result = sb.table("analyses").select("*").eq("id", analysis_id).eq("user_id", user_id).execute()
    if not result.data:
        return None
    row = result.data[0]
    json_fields = ["strengths", "weaknesses", "missing_keywords", "ats_improvements",
                   "formatting_issues", "grammar_issues", "skill_gaps", "recommended_skills",
                   "section_scores", "recommended_certifications", "suggested_projects"]
    for field in json_fields:
        if isinstance(row.get(field), str):
            try:
                row[field] = json.loads(row[field])
            except Exception:
                row[field] = []
    return row


@report_bp.route("/pdf/<analysis_id>", methods=["GET"])
@login_required
def download_pdf(analysis_id):
    user = session["user"]
    a = _get_analysis(analysis_id, user["id"])
    if not a:
        return jsonify({"error": "Analysis not found"}), 404

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import cm

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm,
                                leftMargin=2*cm, rightMargin=2*cm)
        styles = getSampleStyleSheet()
        navy = colors.HexColor("#003580")
        green = colors.HexColor("#16a34a")
        red = colors.HexColor("#dc2626")

        title_style = ParagraphStyle("Title", parent=styles["Heading1"], textColor=navy, fontSize=20)
        h2_style = ParagraphStyle("H2", parent=styles["Heading2"], textColor=navy, fontSize=14)
        body_style = styles["Normal"]

        score = a.get("ats_score", 0)
        score_color = green if score >= 71 else (colors.HexColor("#ca8a04") if score >= 41 else red)

        elements = []
        elements.append(Paragraph("ScoreMyCV — ATS Analysis Report", title_style))
        elements.append(Spacer(1, 0.3*cm))
        elements.append(Paragraph(f"File: {a.get('file_name','N/A')}  |  Grade: {a.get('letter_grade','N/A')}", body_style))
        elements.append(Spacer(1, 0.5*cm))

        # Score box
        score_data = [["ATS Score", f"{score}/100  ({a.get('letter_grade','N/A')})"]]
        score_table = Table(score_data, colWidths=[4*cm, 12*cm])
        score_table.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (0,0), navy),
            ("TEXTCOLOR", (0,0), (0,0), colors.white),
            ("BACKGROUND", (1,0), (1,0), score_color),
            ("TEXTCOLOR", (1,0), (1,0), colors.white),
            ("FONTSIZE", (0,0), (-1,-1), 14),
            ("FONTNAME", (0,0), (-1,-1), "Helvetica-Bold"),
            ("ALIGN", (0,0), (-1,-1), "CENTER"),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0,0), (-1,-1), [None]),
            ("BOX", (0,0), (-1,-1), 1, navy),
            ("TOPPADDING", (0,0), (-1,-1), 10),
            ("BOTTOMPADDING", (0,0), (-1,-1), 10),
        ]))
        elements.append(score_table)
        elements.append(Spacer(1, 0.5*cm))

        if a.get("summary"):
            elements.append(Paragraph("Summary", h2_style))
            elements.append(Paragraph(a["summary"], body_style))
            elements.append(Spacer(1, 0.4*cm))

        def add_list_section(title, items):
            if items:
                elements.append(Paragraph(title, h2_style))
                for item in items:
                    elements.append(Paragraph(f"• {item}", body_style))
                elements.append(Spacer(1, 0.4*cm))

        add_list_section("Strengths", a.get("strengths", []))
        add_list_section("Weaknesses", a.get("weaknesses", []))
        add_list_section("Missing Keywords", a.get("missing_keywords", []))
        add_list_section("ATS Improvement Suggestions", a.get("ats_improvements", []))
        add_list_section("Recommended Skills", a.get("recommended_skills", []))

        if a.get("company_name") or a.get("designation"):
            elements.append(Paragraph("Company / Role Match", h2_style))
            if a.get("company_match_percentage") is not None:
                elements.append(Paragraph(f"Match: {a['company_match_percentage']}%", body_style))
            add_list_section("Recommended Certifications", a.get("recommended_certifications", []))
            add_list_section("Suggested Projects", a.get("suggested_projects", []))

        doc.build(elements)
        buf.seek(0)
        return send_file(buf, mimetype="application/pdf",
                         as_attachment=True,
                         download_name=f"ScoreMyCV_Report_{analysis_id[:8]}.pdf")
    except ImportError:
        return jsonify({"error": "reportlab not installed. Run: pip install reportlab"}), 500


@report_bp.route("/docx/<analysis_id>", methods=["GET"])
@login_required
def download_docx(analysis_id):
    user = session["user"]
    a = _get_analysis(analysis_id, user["id"])
    if not a:
        return jsonify({"error": "Analysis not found"}), 404

    try:
        from docx import Document
        from docx.shared import RGBColor, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        navy = RGBColor(0, 53, 128)

        # Title
        title = doc.add_heading("ScoreMyCV — ATS Analysis Report", 0)
        title.runs[0].font.color.rgb = navy

        doc.add_paragraph(f"File: {a.get('file_name','N/A')}  |  ATS Score: {a.get('ats_score',0)}/100  |  Grade: {a.get('letter_grade','N/A')}")

        def add_section(heading, items):
            if items:
                h = doc.add_heading(heading, level=2)
                h.runs[0].font.color.rgb = navy
                for item in items:
                    doc.add_paragraph(item, style="List Bullet")

        if a.get("summary"):
            h = doc.add_heading("Summary", level=2)
            h.runs[0].font.color.rgb = navy
            doc.add_paragraph(a["summary"])

        add_section("Strengths", a.get("strengths", []))
        add_section("Weaknesses", a.get("weaknesses", []))
        add_section("Missing Keywords", a.get("missing_keywords", []))
        add_section("ATS Improvement Suggestions", a.get("ats_improvements", []))
        add_section("Recommended Skills", a.get("recommended_skills", []))

        if a.get("company_name") or a.get("designation"):
            h = doc.add_heading("Company / Role Match", level=2)
            h.runs[0].font.color.rgb = navy
            if a.get("company_match_percentage") is not None:
                doc.add_paragraph(f"Match Percentage: {a['company_match_percentage']}%")
            add_section("Recommended Certifications", a.get("recommended_certifications", []))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf,
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True,
                         download_name=f"ScoreMyCV_Report_{analysis_id[:8]}.docx")
    except ImportError:
        return jsonify({"error": "python-docx not installed. Run: pip install python-docx"}), 500
